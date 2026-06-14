#!/usr/bin/env python3
"""Tinh chỉnh lại 47 biểu mẫu ISO trong thư mục bm/ theo phản hồi người dùng.

Word (.docx):
  1. Đổi placeholder gạch dưới `____` -> dấu chấm `....` (chỗ điền tay).
  2. Giãn dòng cho các đoạn thân (không áp dụng trong bảng): line spacing 1.5.

Excel (.xlsx):
  1. Đổi placeholder `____` -> `....` ở các trường thông tin/chân biểu (ngoài bảng).
  3. Trong vùng bảng dữ liệu thì để trống, không chèn placeholder (vốn đã trống).
  4. Bỏ cột trống dư thừa ở cuối (mọi merge/nội dung không vượt quá bề rộng bảng).
  5. Khối "Mã số / Ngày ban hành / Lần ban hành" gộp ô vừa đủ (compact bên phải),
     phần tên công ty chiếm phần còn lại bên trái.

Chạy: ~/.claude/skills/.venv/bin/python3 scripts/refine_bm_forms.py
"""
import glob
import os
import re

from docx import Document
from docx.shared import Pt

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

ROOT = os.path.join(os.path.dirname(__file__), "..")
BM = os.path.join(ROOT, "bm")

UNDERSCORE_RUN = re.compile(r"_+")


def dots(text: str) -> str:
    """Đổi mỗi cụm gạch dưới liên tiếp thành cùng số lượng dấu chấm."""
    return UNDERSCORE_RUN.sub(lambda m: "." * len(m.group(0)), text)


# ---------------------------------------------------------------- WORD --------
def fix_docx(path: str) -> str:
    doc = Document(path)
    notes = []

    # (1) placeholder -> dấu chấm trong mọi đoạn (kể cả trong bảng)
    def fix_runs(paragraph):
        for run in paragraph.runs:
            if "_" in run.text:
                run.text = dots(run.text)

    # (2) giãn dòng cho đoạn thân — chỉ các đoạn cấp body, KHÔNG trong bảng
    body_paras = doc.paragraphs  # python-docx chỉ trả đoạn cấp body, bỏ qua bảng
    for p in body_paras:
        fix_runs(p)
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        # đảm bảo có khoảng cách tách đoạn tối thiểu cho dễ đọc
        if pf.space_after is None or pf.space_after < Pt(6):
            pf.space_after = Pt(6)

    # đoạn nằm trong bảng: chỉ đổi placeholder, KHÔNG đổi giãn dòng
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_runs(p)

    doc.save(path)
    return "; ".join(notes) if notes else "dots + line-spacing 1.5"


# --------------------------------------------------------------- EXCEL --------
THIN = Side(style="thin")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
# bề rộng (đơn vị ký tự) tối thiểu cho khối mã số để hiển thị "Ngày ban hành: dd/mm/yyyy"
CODE_BLOCK_MIN_WIDTH = 30.0
DEFAULT_COL_WIDTH = 8.43


def merge_at(ws, row, col):
    for rng in ws.merged_cells.ranges:
        if rng.min_row <= row <= rng.max_row and rng.min_col <= col <= rng.max_col:
            return rng
    return None


def col_width(ws, idx):
    cd = ws.column_dimensions.get(get_column_letter(idx))
    return cd.width if cd and cd.width else DEFAULT_COL_WIDTH


def fix_xlsx(path: str) -> str:
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    notes = []

    # --- xác định kích thước bảng ---
    comp_merge = merge_at(ws, 1, 1)
    header_rows = comp_merge.max_row if comp_merge else 3
    title_row = header_rows + 1
    title_merge = merge_at(ws, title_row, 1)
    ntab = title_merge.max_col if title_merge else ws.max_column

    # --- (5) lấy lại nội dung khối mã số (3 dòng) + tên công ty ---
    company_text = ws.cell(1, 1).value
    code_lines = []
    for r in range(1, header_rows + 1):
        val = None
        for c in range(2, ntab + 1):
            v = ws.cell(r, c).value
            if isinstance(v, str) and v.strip():
                val = v
                break
        code_lines.append(val)
    fill_rgb = ws.cell(1, 1).fill.fgColor.rgb if ws.cell(1, 1).fill.patternType else None

    # bỏ mọi merge trong vùng header (rows 1..header_rows)
    for rng in list(ws.merged_cells.ranges):
        if rng.min_row <= header_rows:
            ws.unmerge_cells(str(rng))
    # xoá nội dung cũ vùng header
    for r in range(1, header_rows + 1):
        for c in range(1, ws.max_column + 1):
            ws.cell(r, c).value = None

    # tính cột bắt đầu khối mã số (gộp vừa đủ từ phải sang)
    acc = 0.0
    cs = ntab
    while cs > 2:
        acc += col_width(ws, cs)
        if acc >= CODE_BLOCK_MIN_WIDTH:
            break
        cs -= 1
    if cs < 2:
        cs = max(2, ntab - 1)

    # tên công ty: A1 : (cs-1, header_rows)
    ws.merge_cells(start_row=1, start_column=1, end_row=header_rows, end_column=cs - 1)
    a1 = ws.cell(1, 1)
    a1.value = company_text
    a1.font = Font(name="Times New Roman", size=13, bold=True)
    a1.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    if fill_rgb:
        a1.fill = PatternFill("solid", fgColor=fill_rgb)

    # khối mã số: mỗi dòng merge cs..ntab
    for i, line in enumerate(code_lines):
        r = i + 1
        ws.merge_cells(start_row=r, start_column=cs, end_row=r, end_column=ntab)
        cell = ws.cell(r, cs)
        cell.value = line
        bold = bool(line and line.strip().lower().startswith("mã số"))
        cell.font = Font(name="Times New Roman", size=11, bold=bold)
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # viền cho toàn bộ ô vùng header
    for r in range(1, header_rows + 1):
        for c in range(1, ntab + 1):
            ws.cell(r, c).border = BORDER

    # --- (1) placeholder gạch -> chấm ở các ô chữ (trường thông tin/chân biểu) ---
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "_" in cell.value:
                cell.value = dots(cell.value)

    # --- (4) kẹp mọi merge về <= ntab rồi xoá cột dư cuối ---
    for rng in list(ws.merged_cells.ranges):
        if rng.max_col > ntab:
            ws.unmerge_cells(str(rng))
            if rng.min_col <= ntab:
                ws.merge_cells(
                    start_row=rng.min_row, start_column=rng.min_col,
                    end_row=rng.max_row, end_column=ntab,
                )
    if ws.max_column > ntab:
        removed = ws.max_column - ntab
        ws.delete_cols(ntab + 1, removed)
        notes.append(f"bỏ {removed} cột dư")

    notes.append(f"header gộp cs={get_column_letter(cs)}..{get_column_letter(ntab)}")
    wb.save(path)
    return "; ".join(notes)


# ----------------------------------------------------------------- MAIN -------
def main():
    docx_files = sorted(glob.glob(os.path.join(BM, "*", "*.docx")))
    xlsx_files = sorted(glob.glob(os.path.join(BM, "*", "*.xlsx")))
    print(f"Word: {len(docx_files)} | Excel: {len(xlsx_files)}\n")
    for f in docx_files:
        note = fix_docx(f)
        print(f"[W] {os.path.basename(f)[:48]:50} {note}")
    print()
    for f in xlsx_files:
        try:
            note = fix_xlsx(f)
            print(f"[X] {os.path.basename(f)[:48]:50} {note}")
        except Exception as e:
            print(f"[X][ERR] {os.path.basename(f)}: {e}")


if __name__ == "__main__":
    main()
